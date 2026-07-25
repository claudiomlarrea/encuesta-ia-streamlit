"""
Informe INSTITUCIONAL extenso (Word), estilo del informe manual del Observatorio.

- Portada estética
- Índice clicable con capítulos y subpuntos
- Numeración de páginas (excepto portada)
- Capítulos: Introducción → tabla → Conclusión
"""
from __future__ import annotations

import io
from collections import defaultdict
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from report_common import (
    CHAPTER_TITLES,
    GREEN,
    MUTED,
    add_aesthetic_cover,
    add_freq_table,
    add_heading,
    add_pandas_table,
    add_para,
    add_toc_hyperlink,
    build_crosstab_section,
    build_frequency_sections,
    build_institutional_final_considerations,
    build_qualitative_extended,
    display_label,
    normalize_user_crosses,
    setup_margins,
)
from survey_intel import classify_columns


CHAPTER_ORDER = [
    "sociodemografico",
    "adopcion",
    "usos",
    "actitudes",
    "institucional",
    "otros",
]


def _infer_audience(source_name: str, columns: list[str]) -> str:
    blob = " ".join([source_name or ""] + columns[:20]).lower()
    if "docente" in blob:
        return "docentes"
    if "alumno" in blob or "estudiante" in blob or "unidad academica estás cursando" in blob:
        return "estudiantes"
    return "la comunidad universitaria"


def _auto_presentation(n: int, n_struct: int, n_open: int, audience: str) -> str:
    return (
        f"El presente informe expone los resultados de la encuesta institucional realizada a "
        f"{audience} de la Universidad Católica de Cuyo sobre el uso, las percepciones y los "
        f"desafíos asociados a la inteligencia artificial. El análisis automático de Encuesta Clara "
        f"procesa {n} respuestas válidas, {n_struct} ítems estructurados"
        + (f" y {n_open} preguntas abiertas" if n_open else "")
        + ". Cada apartado cuantitativo incluye introducción, tabla de frecuencias y conclusión "
        "interpretativa, en la misma lógica del informe institucional elaborado a mano por el "
        "Observatorio. Las lecturas son orientativas y deben validarse en el marco teórico y de gestión."
    )


def _auto_resumen(
    freq_sections: list[dict[str, Any]],
    qual: list[dict[str, Any]],
    n: int,
    *,
    n_crosses: int = 0,
) -> list[str]:
    paras: list[str] = []
    paras.append(
        "La evidencia relevada permite caracterizar patrones de respuesta de la población encuestada "
        f"(N={n}). A continuación se sintetizan los hallazgos de mayor concentración modal y las "
        "lecturas cualitativas emergentes."
    )
    scored = []
    for sec in freq_sections:
        ft = sec["table"]
        body = ft[ft["categoría"].astype(str).str.upper() != "TOTAL"] if "categoría" in ft.columns else ft
        if body.empty or "porcentaje" not in body.columns:
            continue
        scored.append((float(body.iloc[0]["porcentaje"]), sec, str(body.iloc[0]["categoría"])))
    scored.sort(key=lambda x: -x[0])
    bits = []
    for pct, sec, cat in scored[:5]:
        bits.append(f"en «{display_label(sec['label'])}» predomina «{cat}» ({pct:.1f}%)")
    if bits:
        paras.append("Los hallazgos cuantitativos más salientes indican que " + "; ".join(bits) + ".")
    if n_crosses:
        paras.append(
            f"El informe incorpora {n_crosses} lectura(s) desagregada(s) adicional(es), "
            "como subapartados numerados debajo de la pregunta desde la cual se configuraron "
            "(por ejemplo, 1.2.1), con introducción, tabla y conclusión analítica."
        )
    if qual:
        qbits = [f"«{display_label(q['label'])}» ({q['short']})" for q in qual[:2]]
        paras.append(
            "En el plano cualitativo, las respuestas abiertas aportan matices discursivos. "
            + " ".join(qbits)
        )
    paras.append(
        "Se concluye que el diagnóstico automático ofrece una base empírica para orientar "
        "formación, comunicación normativa y acompañamiento pedagógico, siempre con validación "
        "institucional posterior."
    )
    return paras


def _crosses_for_item(crosses: list[dict[str, Any]], column: str | None) -> list[dict[str, Any]]:
    """Solo cruces configurados sobre este ítem (fila = pregunta dueña), sin duplicar en el partner."""
    if not column:
        return []
    return [cr for cr in crosses if cr.get("row_column") == column]


def build_institutional_report_docx(
    df: pd.DataFrame,
    *,
    title: str = "",
    subtitle: str = "",
    source_name: str = "",
    cohort_label: str = "",
    user_crosses: list[dict[str, Any]] | None = None,
    auto_analysis: list[dict[str, Any]] | None = None,
) -> bytes:
    """
    Informe institucional Word.

    Si se pasa `user_crosses` o `auto_analysis` (resultado del Análisis automático),
    el capítulo de cruces usa esos cruces elegidos. Si no, genera cruces automáticos.
    """
    import report_common as rc

    rc._bookmark_seq = 0

    profiles = classify_columns(df)
    freq_sections = build_frequency_sections(df, profiles)
    qual_briefs = build_qualitative_extended(df, profiles)
    if user_crosses is not None:
        crosses = list(user_crosses)
        crosses_from_user = True
    elif auto_analysis is not None:
        crosses = normalize_user_crosses(auto_analysis)
        crosses_from_user = True
    else:
        crosses = build_crosstab_section(df, profiles)
        crosses_from_user = False
    audience = _infer_audience(source_name, [p.name for p in profiles])

    cover_title = title.strip() or (
        f"Diagnóstico institucional sobre el uso de Inteligencia Artificial en {audience} "
        "de la Universidad Católica de Cuyo"
    )
    cover_sub = subtitle.strip() or (
        f"Resultados de la encuesta institucional sobre usos, percepciones y desafíos "
        f"de la Inteligencia Artificial en {audience}."
    )

    by_chapter: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for sec in freq_sections:
        by_chapter[sec["chapter"]].append(sec)

    doc = Document()
    add_aesthetic_cover(
        doc,
        kicker="INFORME INSTITUCIONAL",
        title=cover_title,
        subtitle=cover_sub,
    )
    setup_margins(doc, section=doc.sections[-1])

    # Armar mapa de capítulos / subpuntos antes del índice
    index_map: list[tuple[int, str, str]] = []
    chap_num = 1
    for key in CHAPTER_ORDER:
        if by_chapter.get(key):
            index_map.append((chap_num, key, CHAPTER_TITLES[key]))
            chap_num += 1
    cross_num = None
    # Con cruces del usuario van bajo cada pregunta (p. ej. 1.2.1); no hay capítulo aparte.
    # Con cruces automáticos (CLI / sin plan) sí hay capítulo dedicado.
    if crosses and not crosses_from_user:
        cross_num = chap_num
        chap_num += 1
    qual_num = None
    if qual_briefs:
        qual_num = chap_num

    # ----- ÍNDICE clicable con subpuntos -----
    add_heading(doc, "ÍNDICE", 1, bookmark="bm_indice")
    add_para(
        doc,
        "Hacé clic en cada ítem para ir al apartado correspondiente.",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=6,
    )

    add_toc_hyperlink(doc, "Presentación", "bm_presentacion", bold=True, space_after=3)
    add_toc_hyperlink(doc, "Resumen", "bm_resumen", bold=True, space_after=3)

    for num, key, title_ch in index_map:
        add_toc_hyperlink(doc, f"{num}. {title_ch}", f"bm_ch_{num}", bold=True, space_after=2)
        for i, sec in enumerate(by_chapter[key], start=1):
            label = display_label(sec["label"])
            add_toc_hyperlink(
                doc,
                f"{num}.{i} {label}",
                f"bm_ch_{num}_{i}",
                size=10,
                indent_cm=0.6,
                space_after=1,
            )
            if crosses_from_user:
                for j, cr in enumerate(_crosses_for_item(crosses, sec.get("column")), start=1):
                    add_toc_hyperlink(
                        doc,
                        f"{num}.{i}.{j} {cr['col_label']}",
                        f"bm_ch_{num}_{i}_x{j}",
                        size=9,
                        indent_cm=1.1,
                        space_after=1,
                    )

    cross_chapter_title = "Análisis de cruces entre variables"
    if cross_num is not None:
        add_toc_hyperlink(
            doc,
            f"{cross_num}. {cross_chapter_title}",
            f"bm_ch_{cross_num}",
            bold=True,
            space_after=2,
        )
        for i, cr in enumerate(crosses, start=1):
            add_toc_hyperlink(
                doc,
                f"{cross_num}.{i} {cr['row_label']} × {cr['col_label']}",
                f"bm_ch_{cross_num}_{i}",
                size=10,
                indent_cm=0.6,
                space_after=1,
            )

    if qual_num is not None:
        add_toc_hyperlink(
            doc,
            f"{qual_num}. {CHAPTER_TITLES['cualitativo']}",
            f"bm_ch_{qual_num}",
            bold=True,
            space_after=2,
        )
        for i, b in enumerate(qual_briefs, start=1):
            add_toc_hyperlink(
                doc,
                f"{qual_num}.{i} {display_label(b['label'])}",
                f"bm_ch_{qual_num}_{i}",
                size=10,
                indent_cm=0.6,
                space_after=1,
            )

    add_toc_hyperlink(doc, "Consideraciones finales", "bm_finales", bold=True, space_after=12)

    # ----- Cuerpo -----
    add_heading(doc, "Presentación", 1, bookmark="bm_presentacion")
    add_para(doc, _auto_presentation(len(df), len(freq_sections), len(qual_briefs), audience))
    if cohort_label:
        add_para(doc, f"Cohorte / instrumento: {cohort_label}.", size=10, color=MUTED)

    add_heading(doc, "Resumen", 1, bookmark="bm_resumen")
    for p in _auto_resumen(
        freq_sections,
        qual_briefs,
        len(df),
        n_crosses=len(crosses) if crosses_from_user else 0,
    ):
        add_para(doc, p)

    for num, key, title_ch in index_map:
        sections = by_chapter[key]
        add_heading(doc, f"{num}. {title_ch}", 1, bookmark=f"bm_ch_{num}")
        add_para(
            doc,
            "Este apartado presenta los ítems estructurados detectados automáticamente en esta dimensión. "
            "Cada subapartado incluye introducción, tabla de frecuencias y conclusión interpretativa"
            + (
                "; debajo de cada ítem aparecen los cruces que configuraste en el Análisis automático."
                if crosses_from_user
                else "."
            ),
            size=10,
            color=MUTED,
        )
        for i, sec in enumerate(sections, start=1):
            add_heading(
                doc,
                f"{num}.{i} {display_label(sec['label'])}",
                2,
                bookmark=f"bm_ch_{num}_{i}",
            )
            if sec.get("subtype"):
                add_para(
                    doc,
                    f"Tipo detectado: {sec['subtype']} · N válidos = {sec['n']}",
                    size=9,
                    color=MUTED,
                    space_after=4,
                )
            add_para(doc, "Introducción", size=11, bold=True, color=GREEN, space_after=2)
            add_para(doc, sec["intro"])
            add_freq_table(doc, sec["table"])
            add_para(doc, "Conclusión", size=11, bold=True, color=GREEN, space_after=2)
            add_para(doc, sec["conclusion"])

            if crosses_from_user:
                item_crosses = _crosses_for_item(crosses, sec.get("column"))
                for j, cr in enumerate(item_crosses, start=1):
                    add_heading(
                        doc,
                        f"{num}.{i}.{j} {cr['col_label']}",
                        3,
                        bookmark=f"bm_ch_{num}_{i}_x{j}",
                    )
                    add_para(doc, "Introducción", size=11, bold=True, color=GREEN, space_after=2)
                    add_para(doc, cr["intro"])
                    add_pandas_table(doc, cr["table"])
                    add_para(doc, "Conclusión", size=11, bold=True, color=GREEN, space_after=2)
                    add_para(doc, cr["conclusion"])

    if crosses and cross_num is not None:
        add_heading(
            doc,
            f"{cross_num}. {cross_chapter_title}",
            1,
            bookmark=f"bm_ch_{cross_num}",
        )
        add_para(
            doc,
            "Se incluyen cruces automáticos entre una variable sociodemográfica y variables de "
            "adopción/uso/institucionales, para explorar diferencias de patrón según perfil.",
        )
        for i, cr in enumerate(crosses, start=1):
            add_heading(
                doc,
                f"{cross_num}.{i} {cr['row_label']} × {cr['col_label']}",
                2,
                bookmark=f"bm_ch_{cross_num}_{i}",
            )
            add_para(doc, "Introducción", size=11, bold=True, color=GREEN, space_after=2)
            add_para(doc, cr["intro"])
            add_pandas_table(doc, cr["table"])
            add_para(doc, "Conclusión", size=11, bold=True, color=GREEN, space_after=2)
            add_para(doc, cr["conclusion"])

    if qual_briefs and qual_num is not None:
        add_heading(
            doc,
            f"{qual_num}. {CHAPTER_TITLES['cualitativo']}",
            1,
            bookmark=f"bm_ch_{qual_num}",
        )
        add_para(
            doc,
            "Análisis temático, de sentimiento y del discurso sobre las respuestas abiertas. "
            "Los resultados son exploratorios y orientativos.",
        )
        for i, b in enumerate(qual_briefs, start=1):
            add_heading(
                doc,
                f"{qual_num}.{i} {display_label(b['label'])}",
                2,
                bookmark=f"bm_ch_{qual_num}_{i}",
            )
            add_para(doc, f"N = {b['n']} respuestas válidas", size=9, color=MUTED)
            add_para(doc, "Síntesis", size=11, bold=True, color=GREEN, space_after=2)
            add_para(doc, b["short"])

            add_heading(doc, f"{qual_num}.{i}.1 Análisis temático", 3)
            for para in b["thematic_paras"][:36]:
                add_para(doc, para, size=10, space_after=5)

            add_heading(doc, f"{qual_num}.{i}.2 Análisis de sentimiento", 3)
            for para in b["sentiment_paras"][:28]:
                add_para(doc, para, size=10, space_after=5)

            add_heading(doc, f"{qual_num}.{i}.3 Análisis del discurso", 3)
            for para in b["discourse_paras"][:36]:
                add_para(doc, para, size=10, space_after=5)

    add_heading(doc, "Consideraciones finales", 1, bookmark="bm_finales")
    add_para(
        doc,
        build_institutional_final_considerations(
            len(df), len(freq_sections), len(qual_briefs), audience
        ),
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
