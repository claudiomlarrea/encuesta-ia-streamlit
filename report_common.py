"""
Utilidades compartidas para informes Word (ejecutivo e institucional).
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from qualitative_deep import (
    deep_discourse_markdown,
    deep_sentiment_markdown,
    deep_thematic_markdown,
)
from quant_advanced import descriptive_one_column
from survey_intel import (
    ColumnProfile,
    add_total_count_row,
    classify_columns,
    frequency_table,
    kwic_snippets,
    lexicon_sentiment_es,
    ngram_top_table,
    thematic_nmf,
)

GREEN = RGBColor(0x04, 0x4A, 0x30)
MAROON = RGBColor(0x7A, 0x15, 0x32)
TEXT = RGBColor(0x1A, 0x2E, 0x28)
MUTED = RGBColor(0x5C, 0x4F, 0x54)
LOGO_PATH = Path(__file__).resolve().parent / "assets" / "logo_observatorio_ia.png"

TEAM_BLOCK = [
    ("Dirección general", "Claudio Larrea Arnau\nDirector del Observatorio de Inteligencia Artificial"),
    (
        "Equipo ejecutivo",
        "Belén Arias\nJavier Coria\nJosé La Malfa\nLaura Pizarro\nStefania Young",
    ),
    ("Asesor externo", "Frederic Marimon\nUniversitat Internacional de Catalunya"),
    ("Elaboración técnica y análisis", "Observatorio de Inteligencia Artificial — UCCuyo"),
]


def short_label(col: str, max_len: int | None = 110) -> str:
    s = re.sub(r"\s+", " ", str(col)).strip()
    m = re.search(r"\[([^\]]+)\]\s*$", s)
    if m:
        s = m.group(1).strip()
    if max_len is not None and len(s) > max_len:
        return s[: max_len - 1] + "…"
    return s


def display_label(col: str) -> str:
    """Etiqueta completa sin recortar (para títulos y conclusiones del institucional)."""
    return short_label(col, max_len=None)


def limit_words(text: str, max_words: int = 60) -> str:
    words = re.findall(r"\S+", (text or "").strip())
    if len(words) <= max_words:
        return " ".join(words)
    return " ".join(words[:max_words]).rstrip(",.;:") + "."


def ensure_word_count(text: str, *, min_words: int = 55, max_words: int = 60, pad: str = "") -> str:
    """Asegura interpretaciones cercanas a 60 palabras (ni muy escuetas ni excesivas)."""
    base = re.sub(r"\s+", " ", (text or "").strip())
    words = re.findall(r"\S+", base)
    if len(words) > max_words:
        return " ".join(words[:max_words]).rstrip(",.;:") + "."
    if len(words) >= min_words:
        return base
    filler = pad or (
        "Esta lectura descriptiva orienta la priorización institucional y debe complementarse "
        "con cruces demográficos, el contexto de cada unidad académica y la validación experta "
        "antes de traducir el hallazgo en decisiones de política universitaria."
    )
    merged = (base.rstrip(".") + ". " + filler).strip()
    return limit_words(merged, max_words)

def set_run(run, *, size: int = 11, bold: bool = False, color: RGBColor | None = None, font: str = "Calibri") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rPr.append(rFonts)


def add_para(
    doc: Document,
    text: str,
    *,
    size: int = 11,
    bold: bool = False,
    color: RGBColor | None = None,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_after: float = 8,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color or TEXT)


def add_heading(doc: Document, text: str, level: int = 1, *, bookmark: str | None = None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run(run, size=16 if level == 1 else (13 if level == 2 else 11), bold=True, color=MAROON if level <= 2 else GREEN)
    if bookmark:
        add_bookmark(h, bookmark)
    return h


_bookmark_seq = 0


def _next_bookmark_id() -> int:
    global _bookmark_seq
    _bookmark_seq += 1
    return _bookmark_seq


def sanitize_bookmark_name(name: str) -> str:
    """Nombre válido de bookmark Word (sin espacios; empieza con letra)."""
    raw = re.sub(r"[^A-Za-z0-9_]", "_", (name or "bm").strip())
    raw = re.sub(r"_+", "_", raw).strip("_")
    if not raw:
        raw = "bm"
    if raw[0].isdigit():
        raw = "bm_" + raw
    return raw[:40]


def add_bookmark(paragraph, name: str) -> str:
    """Inserta bookmark alrededor del contenido del párrafo. Devuelve el nombre usado."""
    bm = sanitize_bookmark_name(name)
    bid = str(_next_bookmark_id())
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), bm)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)
    return bm


def add_toc_hyperlink(
    doc: Document,
    text: str,
    bookmark: str,
    *,
    size: int = 11,
    bold: bool = False,
    indent_cm: float = 0.0,
    space_after: float = 2,
) -> None:
    """Entrada de índice clicable hacia un bookmark interno."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)

    bm = sanitize_bookmark_name(bookmark)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bm)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)
    if bold:
        rPr.append(OxmlElement("w:b"))
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)


def add_page_numbers_bottom_right(section, *, start_at: int = 1) -> None:
    """Numeración de página abajo a la derecha (solo esta sección; no en portada)."""
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    # Vaciar contenido previo del párrafo
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    def _run_with(*children):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")
        rPr.append(sz)
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "5C4F54")
        rPr.append(color_el)
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(rFonts)
        r.append(rPr)
        for ch in children:
            r.append(ch)
        return r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = str(start_at)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    p._p.append(_run_with(fld_begin))
    p._p.append(_run_with(instr))
    p._p.append(_run_with(fld_sep))
    p._p.append(_run_with(fld_text))
    p._p.append(_run_with(fld_end))

    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start_at))


def shade_paragraph(paragraph, hex_fill: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def add_freq_table(doc: Document, ft: pd.DataFrame, max_rows: int = 20) -> None:
    cols = [c for c in ("categoría", "frecuencia", "porcentaje") if c in ft.columns]
    if not cols:
        cols = list(ft.columns)[:3]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    labels = {"categoría": "Categoría", "frecuencia": "Frecuencia", "porcentaje": "Porcentaje"}
    for j, c in enumerate(cols):
        hdr[j].text = labels.get(c, str(c))
        for p in hdr[j].paragraphs:
            for run in p.runs:
                set_run(run, size=10, bold=True, color=GREEN)
    body = ft.copy()
    if "categoría" in body.columns:
        body = body[body["categoría"].astype(str).str.upper() != "TOTAL"]
    for _, row in body.head(max_rows).iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(cols):
            val = row.get(c, "")
            if c == "porcentaje" and isinstance(val, (int, float)):
                cells[j].text = f"{float(val):.2f}"
            else:
                cells[j].text = str(val)
            for p in cells[j].paragraphs:
                for run in p.runs:
                    set_run(run, size=10, color=TEXT)
    doc.add_paragraph("")


def md_to_plain_paragraphs(md: str) -> list[str]:
    out: list[str] = []
    buf: list[str] = []
    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            if buf:
                out.append(" ".join(buf))
                buf = []
            continue
        if line.lstrip().startswith("|") or set(line.strip()) <= {"-", "|", ":", " "}:
            continue
        if line.startswith("#"):
            if buf:
                out.append(" ".join(buf))
                buf = []
            title = re.sub(r"^#+\s*", "", line).strip()
            title = re.sub(r"\*+", "", title)
            if title:
                out.append(title)
            continue
        if line.lstrip().startswith(("-", "*", "•")):
            if buf:
                out.append(" ".join(buf))
                buf = []
            item = re.sub(r"^[\-\*•]\s*", "", line.strip())
            item = re.sub(r"\*+", "", item)
            item = re.sub(r"`+", "", item)
            out.append("• " + item)
            continue
        clean = re.sub(r"\*+", "", line)
        clean = re.sub(r"`+", "", clean)
        clean = re.sub(r"<[^>]+>", "", clean)
        buf.append(clean.strip())
    if buf:
        out.append(" ".join(buf))
    return [p for p in out if p and len(p.strip()) > 2]


def top_category(ft: pd.DataFrame) -> tuple[str, float, int]:
    body = ft.copy()
    if "categoría" in body.columns:
        body = body[body["categoría"].astype(str).str.upper() != "TOTAL"]
    if body.empty:
        return "", 0.0, 0
    row = body.iloc[0]
    cat = str(row.get("categoría", row.iloc[0]))
    pct = float(row["porcentaje"]) if "porcentaje" in body.columns else 0.0
    n = int(row["frecuencia"]) if "frecuencia" in body.columns else 0
    return cat, pct, n


def interpret_frequency_short(label: str, ft: pd.DataFrame, n_valid: int, subtype: str = "") -> str:
    if ft is None or ft.empty or n_valid <= 0:
        return ensure_word_count("Sin respuestas válidas suficientes para interpretar este ítem.")
    cat, pct, n_cat = top_category(ft)
    body = ft[ft["categoría"].astype(str).str.upper() != "TOTAL"] if "categoría" in ft.columns else ft
    st = (subtype or "").lower()
    lab = display_label(label)
    second = ""
    if len(body) > 1 and "porcentaje" in body.columns:
        c2 = str(body.iloc[1].get("categoría", ""))
        p2 = float(body.iloc[1]["porcentaje"])
        n2 = int(body.iloc[1]["frecuencia"]) if "frecuencia" in body.columns else 0
        second = f" La segunda categoría es «{c2}» ({p2:.1f}%, n={n2})."

    if "likert" in st or "acuerdo" in st:
        base = (
            f"En «{lab}», predomina «{cat}» ({pct:.1f}%, n={n_cat}) sobre {n_valid} respuestas válidas."
            f"{second} La distribución orienta la lectura de adhesión o reserva frente al enunciado "
            f"y permite anticipar tensiones pedagógicas o de gobernanza asociadas al ítem."
        )
    elif "frecuencia" in st:
        base = (
            f"En «{lab}», el uso declarado se concentra en «{cat}» ({pct:.1f}%, n={n_cat}) "
            f"sobre un total de {n_valid} respuestas."
            f"{second} "
        )
        if pct >= 60:
            base += "El patrón indica una práctica extendida o recurrente en la muestra analizada."
        elif pct <= 25:
            base += "Sugiere heterogeneidad: ninguna categoría concentra a la mayoría de respondentes."
        else:
            base += "Hay un polo dominante, con margen de variación relevante entre respondentes."
        base += " Conviene leer este corte junto con demografía y con la dimensión normativa institucional."
    elif "binaria" in st:
        base = (
            f"En «{lab}», la opción mayoritaria es «{cat}» ({pct:.1f}%, n={n_cat}) de {n_valid}."
            f"{second} Este corte resulta útil para estimar penetración o presencia/ausencia del fenómeno "
            f"y para orientar acciones de comunicación o formación diferenciadas."
        )
    else:
        base = (
            f"En «{lab}», la categoría más frecuente es «{cat}» ({pct:.1f}%, n={n_cat}) entre "
            f"{n_valid} respuestas válidas y {len(body)} categorías observadas."
            f"{second} El perfil modal guía el diagnóstico descriptivo y la comparación entre unidades académicas."
        )
    return ensure_word_count(base, min_words=55, max_words=60)


def interpret_frequency_long(label: str, ft: pd.DataFrame, n_valid: int, subtype: str = "") -> tuple[str, str]:
    """Devuelve (introducción, conclusión) estilo informe institucional — sin cortar etiquetas."""
    lab = display_label(label)
    cat, pct, n_cat = top_category(ft)
    body = ft[ft["categoría"].astype(str).str.upper() != "TOTAL"] if "categoría" in ft.columns else ft
    n_cats = len(body)

    intro = (
        f"La presente tabla describe la distribución de respuestas al ítem «{lab}». "
        f"El análisis permite caracterizar el patrón predominante en la muestra (N={n_valid}) "
        f"y orientar la lectura institucional del fenómeno relevado, en diálogo con el resto del capítulo "
        f"y con los cruces sociodemográficos."
    )

    details = []
    for _, row in body.head(5).iterrows():
        c = str(row.get("categoría", ""))
        f = int(row["frecuencia"]) if "frecuencia" in body.columns else 0
        p = float(row["porcentaje"]) if "porcentaje" in body.columns else 0.0
        details.append(f"«{c}» ({p:.2f}%, n={f})")
    detail_txt = "; ".join(details)

    st = (subtype or "").lower()
    if "likert" in st or "acuerdo" in st:
        concl = (
            f"Los resultados muestran que la posición modal es «{cat}» ({pct:.2f}%, {n_cat} casos) "
            f"sobre un total de {n_valid} respuestas válidas. "
            f"Las categorías más frecuentes son: {detail_txt}. "
            f"Esta distribución informa el grado de adhesión o reserva frente al enunciado y "
            f"debe leerse junto con los cruces demográficos y los discursos abiertos."
        )
    elif "frecuencia" in st:
        concl = (
            f"Se observa una concentración en «{cat}» ({pct:.2f}%, {n_cat} casos). "
            f"El detalle de las principales categorías es: {detail_txt}. "
            f"El patrón sugiere el grado de integración de la práctica relevada en la rutina "
            f"de la población encuestada (N={n_valid}) y aporta insumos para priorizar acompañamiento pedagógico."
        )
    elif "binaria" in st:
        concl = (
            f"La opción mayoritaria es «{cat}» ({pct:.2f}%, {n_cat} casos) entre {n_valid} respondentes. "
            f"Distribución principal: {detail_txt}. "
            f"Este corte binario resulta útil para estimar penetración o presencia/ausencia del fenómeno "
            f"y para diseñar intervenciones de comunicación institucional."
        )
    else:
        concl = (
            f"La categoría modal es «{cat}» ({pct:.2f}%, {n_cat} casos) "
            f"entre {n_valid} respuestas y {n_cats} categorías observadas. "
            f"Principales valores: {detail_txt}. "
            f"La lectura debe considerar la heterogeneidad de etiquetas y el eventual carácter "
            f"multirrespuesta del ítem, así como diferencias por unidad académica."
        )
    return intro, concl

def classify_chapter(profile: ColumnProfile) -> str:
    """Agrupa ítems en capítulos del informe institucional."""
    name = profile.name.lower()
    st = (profile.subtype or "").lower()
    if profile.kind == "abierta":
        return "cualitativo"

    # Grillas de usos (Google Forms) — priorizar antes de demografía
    if (
        "usos posibles" in name
        or "indicá con qué frecuencia" in name
        or "indica con que frecuencia" in name
        or ("[" in profile.name and "frecuencia" in name)
    ):
        return "usos"

    inst_keys = (
        "normativa",
        "advert",
        "capacita",
        "formación sobre",
        "formacion sobre",
        "temas te gustaría capacitar",
        "declaración",
        "transparencia",
        "postura predominante",
        "trabajos completos sin intervención",
        "falta?",
        "debería considerarse una falta",
    )
    if any(k in name for k in inst_keys):
        return "institucional"

    adopt_keys = (
        "conocés o usaste",
        "conoces o usaste",
        "herramientas de ia conocés",
        "herramientas de ia conoces",
        "frecuencia utilizás herramientas",
        "frecuencia utilizas herramientas",
        "dónde conociste",
        "donde conociste",
        "alguna vez una herramienta de inteligencia",
    )
    if any(k in name for k in adopt_keys):
        return "adopcion"

    if "likert" in st or "afirmaciones" in name or "nivel de acuerdo" in name:
        return "actitudes"

    dem_keys = (
        "unidad academica estás cursando",
        "unidad académica estás cursando",
        "unidad academica dictás",
        "unidad académica dictás",
        "unidad academica dictas",
        "año de la carrera",
        "ano de la carrera",
        "edad:",
        "género:",
        "genero:",
        "trabajás actualmente",
        "trabajas actualmente",
        "acceso frecuente a una computadora",
        "sede principal",
        "categoría / vínculo",
        "categoria / vinculo",
        "antigüedad en la docencia",
        "antiguedad en la docencia",
    )
    if any(k in name for k in dem_keys):
        return "sociodemografico"
    # demografía corta típica
    if name.strip() in {
        "edad",
        "edad:",
        "género",
        "género:",
        "genero",
        "genero:",
    }:
        return "sociodemografico"

    if "frecuencia" in st:
        return "usos"
    return "otros"


CHAPTER_TITLES = {
    "sociodemografico": "Caracterización sociodemográfica y acceso digital",
    "adopcion": "Nivel de conocimiento, adopción y frecuencia de uso",
    "usos": "Usos y prácticas declaradas",
    "actitudes": "Percepciones, impacto pedagógico y actitudes",
    "institucional": "Dimensión institucional, regulación y formación",
    "cualitativo": "Representaciones sociales y discursos (análisis cualitativo)",
    "otros": "Otros hallazgos cuantitativos",
}


def build_frequency_sections(df: pd.DataFrame, profiles: list[ColumnProfile]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    for p in profiles:
        if p.kind != "estructurada" or p.n_non_null <= 0:
            continue
        try:
            ft = frequency_table(df[p.name])
            desc = descriptive_one_column(df[p.name])
        except Exception:  # noqa: BLE001
            continue
        if ft is None or ft.empty:
            continue
        # Etiqueta completa (texto entre [ ] si es grilla Forms), sin recortar
        label = display_label(p.name)
        intro, concl = interpret_frequency_long(label, ft, int(p.n_non_null), p.subtype)
        sections.append(
            {
                "column": p.name,
                "label": label,
                "subtype": p.subtype,
                "table": ft,
                "desc": desc,
                "n": int(p.n_non_null),
                "chapter": classify_chapter(p),
                "intro": intro,
                "conclusion": concl,
                "short": interpret_frequency_short(label, ft, int(p.n_non_null), p.subtype),
            }
        )
    return sections


def build_qualitative_extended(df: pd.DataFrame, profiles: list[ColumnProfile], n_topics: int = 5) -> list[dict[str, Any]]:
    briefs: list[dict[str, Any]] = []
    open_cols = [p for p in profiles if p.kind == "abierta" and p.n_non_null >= 8]
    for p in open_cols[:8]:
        col = p.name
        label = display_label(col)
        texts = [t.strip() for t in df[col].dropna().astype(str).tolist() if len(t.strip()) > 12]
        if len(texts) < 8:
            continue
        topics, _W, dominant, quotes, texts_nmf = [], [], [], {}, []
        try:
            k = min(n_topics, max(2, len(texts) // 12))
            topics, _W, dominant, quotes, texts_nmf = thematic_nmf(texts, n_topics=k)
        except Exception:  # noqa: BLE001
            pass
        thematic_md = deep_thematic_markdown(label, topics, dominant, quotes, texts_nmf, corpus=texts)

        results: list[str] = []
        filtered: list[str] = []
        for t in texts:
            lab, _score = lexicon_sentiment_es(t)
            results.append(lab)
            filtered.append(t)
        dist = pd.Series(results).value_counts().rename_axis("sentimiento").reset_index(name="n")
        dist["pct"] = (dist["n"] / max(len(results), 1) * 100).round(1)
        dist = add_total_count_row(dist, label_col="sentimiento", value_col="n")
        dist.loc[dist["sentimiento"] == "TOTAL", "pct"] = 100.0
        sentiment_md = deep_sentiment_markdown(
            label, filtered, results, dist, metodo="Léxico en español (orientativo)"
        )

        bi = ngram_top_table(texts, ngram_range=(2, 2), top_n=20)
        tri = ngram_top_table(texts, ngram_range=(3, 3), top_n=15)
        needle = ""
        hits: list[str] = []
        for candidate in ("inteligencia", "riesgo", "formación", "ética", "aprendizaje", "docente"):
            cand_hits = kwic_snippets(texts, candidate, max_hits=6, half_window=60)
            if cand_hits:
                needle = candidate
                hits = cand_hits
                break
        discourse_md = deep_discourse_markdown(label, texts, bi, tri, needle, hits)

        pos = sum(1 for r in results if r == "positivo")
        neu = sum(1 for r in results if r == "neutral")
        neg = sum(1 for r in results if r == "negativo")
        total = max(len(results), 1)
        top_theme = ""
        if topics:
            top_theme = str(topics[0].get("palabras_clave", "")).split(",")[0].strip()
        short = (
            f"En «{display_label(label)}» (N={len(texts)}), el tono se distribuye "
            f"positivo {pos/total*100:.0f}%, neutral {neu/total*100:.0f}% y negativo {neg/total*100:.0f}%. "
        )
        if top_theme:
            short += f"El eje temático más saliente asocia «{top_theme}». "
        else:
            short += "El corpus abre lecturas sobre usos, riesgos y expectativas formativas. "
        short += (
            "Esta síntesis cualitativa es exploratoria y debe validarse con lectura fina del corpus "
            "antes de su uso en decisiones institucionales."
        )

        briefs.append(
            {
                "label": label,
                "n": len(texts),
                "short": ensure_word_count(short, min_words=55, max_words=60),
                "thematic_paras": md_to_plain_paragraphs(thematic_md),
                "sentiment_paras": md_to_plain_paragraphs(sentiment_md),
                "discourse_paras": md_to_plain_paragraphs(discourse_md),
            }
        )
    return briefs


def _dem_priority(name: str) -> int:
    n = name.lower()
    if "año de la carrera" in n or "ano de la carrera" in n:
        return 0
    if "unidad academica" in n or "unidad académica" in n:
        return 1
    if "género" in n or "genero" in n:
        return 2
    if "edad" in n:
        return 3
    return 9


def _target_priority(name: str) -> int:
    n = name.lower()
    if "computadora" in n or "notebook" in n:
        return 0
    if "trabajás actualmente" in n or "trabajas actualmente" in n:
        return 1
    if "conocés o usaste" in n or "conoces o usaste" in n:
        return 2
    if "frecuencia utiliz" in n:
        return 3
    if "normativa" in n or "capacita" in n or "advert" in n:
        return 4
    return 8


def normalize_user_crosses(computed: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    """
    Convierte el resultado de Análisis automático (`auto_cross_computed`)
    en la lista de cruces que consumen los informes Word.
    """
    out: list[dict[str, Any]] = []
    if not computed:
        return out
    for item in computed:
        row_label = display_label(str(item.get("label") or item.get("column") or ""))
        row_column = item.get("column")
        for cr in item.get("crosses") or []:
            table = cr.get("table")
            if table is None:
                continue
            try:
                if getattr(table, "empty", False) or table.shape[0] < 1 or table.shape[1] < 1:
                    continue
            except Exception:  # noqa: BLE001
                continue
            col_label = display_label(str(cr.get("partner_label") or cr.get("partner") or ""))
            out.append(
                {
                    "row_label": row_label,
                    "col_label": col_label,
                    "row_column": row_column,
                    "col_column": cr.get("partner"),
                    "table": table,
                    "intro": (
                        f"El cruce entre «{row_label}» y «{col_label}» fue seleccionado "
                        "en el Análisis automático para explorar la distribución conjunta "
                        "y posibles diferencias de patrón entre subgrupos."
                    ),
                    "conclusion": (
                        f"La tabla de contingencia ({table.shape[0]} filas × {table.shape[1]} columnas) "
                        "complementa la lectura univariada. Cuando el tamaño muestral lo permita, "
                        "conviene contrastar con una prueba de asociación (χ²) en Encuesta Clara."
                    ),
                }
            )
    return out


def narrative_for_crosstab(cr: dict[str, Any], *, executive: bool = False) -> str:
    """Párrafo narrativo a partir de una tabla de contingencia."""
    table = cr.get("table")
    row_label = cr.get("row_label", "")
    col_label = cr.get("col_label", "")
    if table is None:
        if executive:
            return (
                f"La lectura conjunta de «{row_label}» y «{col_label}» no pudo completarse "
                "por falta de datos suficientes en esta corrida."
            )
        return (
            f"Se consideró el cruce «{row_label}» × «{col_label}», sin tabla disponible."
        )
    try:
        flat = table.stack()
        if flat.empty:
            raise ValueError("empty")
        idx = flat.idxmax()
        if isinstance(idx, tuple) and len(idx) == 2:
            r_lab, c_lab = idx
        else:
            r_lab, c_lab = str(idx), ""
        val = int(flat.max())
        total = int(flat.sum())
        pct = (100.0 * val / total) if total else 0.0
        r_disp = display_label(str(r_lab))
        c_disp = display_label(str(c_lab))
        if executive:
            return (
                f"Al considerar de manera conjunta «{row_label}» y «{col_label}», el patrón más frecuente "
                f"asocia «{r_disp}» con «{c_disp}» ({val} casos; {pct:.1f}% del total observado, N={total}). "
                "Esta lectura bivariada aporta matices según perfil y debe interpretarse como señal "
                "exploratoria para la conducción académica, no como prueba causal."
            )
        return (
            f"En el cruce «{row_label}» × «{col_label}», la celda más frecuente asocia "
            f"«{r_disp}» con «{c_disp}» "
            f"({val} casos; {pct:.1f}% del total cruzado, N={total}). "
            "La lectura es exploratoria y debe validarse con el contexto institucional."
        )
    except Exception:  # noqa: BLE001
        if executive:
            return (
                f"La lectura conjunta de «{row_label}» y «{col_label}» aporta matices de distribución "
                "según subgrupos y complementa los hallazgos univariados del diagnóstico."
            )
        return (
            f"El cruce «{row_label}» × «{col_label}» aporta la distribución conjunta "
            f"({getattr(table, 'shape', ('?', '?'))[0]}×{getattr(table, 'shape', ('?', '?'))[1]}). "
            "Complementa el análisis univariado del informe."
        )


def build_crosstab_section(df: pd.DataFrame, profiles: list[ColumnProfile]) -> list[dict[str, Any]]:
    """Cruces automáticos demográficos × adopción/uso/acceso (estilo informe manual)."""
    dem = sorted(
        [p for p in profiles if classify_chapter(p) == "sociodemografico" and p.kind == "estructurada"],
        key=lambda p: _dem_priority(p.name),
    )
    # Incluye demografía “acceso/trabajo” como columnas (como en el manual 1.7 y 1.8)
    dem_as_target = [
        p
        for p in profiles
        if classify_chapter(p) == "sociodemografico"
        and p.kind == "estructurada"
        and any(k in p.name.lower() for k in ("computadora", "notebook", "trabajás actualmente", "trabajas actualmente"))
    ]
    targets = sorted(
        [
            p
            for p in profiles
            if (
                classify_chapter(p) in {"adopcion", "usos", "institucional"}
                or p in dem_as_target
            )
            and p.kind == "estructurada"
        ],
        key=lambda p: _target_priority(p.name),
    )
    out: list[dict[str, Any]] = []
    if not dem or not targets:
        return out

    pairs: list[tuple[ColumnProfile, ColumnProfile]] = []
    # Preferir año × acceso / año × trabajo (como el informe manual)
    yearish = [p for p in dem if _dem_priority(p.name) == 0]
    if yearish:
        for t in targets:
            if t.name == yearish[0].name:
                continue
            pairs.append((yearish[0], t))
            if len(pairs) >= 2:
                break
    # Luego unidad × adopción / institucional
    ua = [p for p in dem if _dem_priority(p.name) == 1]
    if ua:
        for t in targets:
            if classify_chapter(t) not in {"adopcion", "institucional"}:
                continue
            if t.name == ua[0].name:
                continue
            pairs.append((ua[0], t))
            if len(pairs) >= 4:
                break
    # Completar si faltan
    for row_p in dem[:2]:
        for col_p in targets:
            if row_p.name == col_p.name:
                continue
            if (row_p, col_p) in pairs or any(a.name == row_p.name and b.name == col_p.name for a, b in pairs):
                continue
            pairs.append((row_p, col_p))
            if len(pairs) >= 4:
                break
        if len(pairs) >= 4:
            break

    for row_p, col_p in pairs[:4]:
        try:
            ct = pd.crosstab(df[row_p.name], df[col_p.name])
            if ct.shape[0] < 2 or ct.shape[1] < 2:
                continue
            if ct.shape[0] > 12:
                ct = ct.iloc[:12]
            if ct.shape[1] > 8:
                ct = ct.iloc[:, :8]
            out.append(
                {
                    "row_label": display_label(row_p.name),
                    "col_label": display_label(col_p.name),
                    "table": ct,
                    "intro": (
                        f"El análisis de la relación entre «{display_label(row_p.name)}» y "
                        f"«{display_label(col_p.name)}» permite explorar posibles diferencias de "
                        "patrón según el perfil de la población encuestada."
                    ),
                    "conclusion": (
                        f"La tabla de contingencia muestra la distribución conjunta "
                        f"(filas={ct.shape[0]}, columnas={ct.shape[1]}). "
                        "La lectura complementa el análisis univariado; cuando corresponda, "
                        "conviene contrastar con pruebas de asociación (χ²) en Encuesta Clara."
                    ),
                }
            )
        except Exception:  # noqa: BLE001
            continue
    return out


def add_pandas_table(doc: Document, table_df: pd.DataFrame) -> None:
    cols = [str(c) for c in table_df.columns]
    index_name = str(table_df.index.name or "Categoría")
    t = doc.add_table(rows=1, cols=len(cols) + 1)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = index_name
    for j, c in enumerate(cols):
        hdr[j + 1].text = display_label(c)
    for idx, row in table_df.iterrows():
        cells = t.add_row().cells
        cells[0].text = display_label(str(idx))
        for j, c in enumerate(cols):
            cells[j + 1].text = str(int(row[c])) if pd.notna(row[c]) else "0"
    doc.add_paragraph("")


def setup_margins(doc: Document, section=None) -> None:
    sections = [section] if section is not None else list(doc.sections)
    for sec in sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)


def add_aesthetic_cover(
    doc: Document,
    *,
    kicker: str,
    title: str,
    subtitle: str,
) -> None:
    """Portada a página completa (PNG) con estética bordo/verde institucional."""
    from cover_render import render_cover_png

    png = render_cover_png(kicker=kicker, title=title, subtitle=subtitle)
    sec0 = doc.sections[0]
    sec0.page_width = Cm(21.0)
    sec0.page_height = Cm(29.7)
    sec0.top_margin = Cm(0)
    sec0.bottom_margin = Cm(0)
    sec0.left_margin = Cm(0)
    sec0.right_margin = Cm(0)
    sec0.header_distance = Cm(0)
    sec0.footer_distance = Cm(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    pf = p.paragraph_format
    try:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    except Exception:  # noqa: BLE001
        pass
    run = p.add_run()
    # Forzar A4 completo (ancho y alto) para que la tipografía de la imagen se lea grande
    run.add_picture(io.BytesIO(png), width=Cm(21.0), height=Cm(29.7))

    # Cuerpo en nueva sección con márgenes normales
    new_sec = doc.add_section()
    setup_margins(doc, section=new_sec)
    # Portada sin número; cuerpo con numeración abajo a la derecha
    try:
        sec0.footer.is_linked_to_previous = False
        for p in sec0.footer.paragraphs:
            p.clear()
    except Exception:  # noqa: BLE001
        pass
    add_page_numbers_bottom_right(new_sec, start_at=1)


def build_institutional_final_considerations(n: int, n_struct: int, n_open: int, audience: str) -> str:
    """Consideraciones finales ~200 palabras para el informe institucional."""
    text = (
        f"El informe institucional automático del Observatorio de Inteligencia Artificial ofrece una base "
        f"empírica para la discusión académica y de gestión a partir de {n} respuestas de {audience}, "
        f"{n_struct} ítems estructurados"
        + (f" y {n_open} preguntas abiertas" if n_open else "")
        + ". Las frecuencias, los cruces sociodemográficos y las lecturas cualitativas (temáticas, de "
        "sentimiento y de discurso) deben contrastarse con el contexto de cada unidad académica, con el "
        "marco normativo vigente y con criterios éticos de uso de datos personales. "
        "La evidencia relevada confirma que la inteligencia artificial ya forma parte de las prácticas "
        "cotidianas de estudio y trabajo universitario: no alcanza con observar el fenómeno; corresponde "
        "gobernarlo con criterios pedagógicos, de integridad académica y de equidad entre sedes y carreras. "
        "En ese sentido, los hallazgos orientan prioridades de formación estudiantil y docente, de "
        "comunicación normativa visible, de revisión de consignas y evaluaciones, y de monitoreo periódico "
        "de usos y percepciones. "
        "Las interpretaciones automáticas son exploratorias y no reemplazan la validación experta ni el "
        "debate colegiado entre Rectorado, Consejo Superior y unidades académicas. "
        "Se recomienda utilizar este documento como insumo de trabajo institucional —complementado con "
        "lectura fina de tablas y corpus abiertos— para construir una política universitaria coherente que "
        "aproveche el potencial formativo de la IA sin renunciar al pensamiento crítico, la autoría y la "
        "centralidad del conocimiento. "
        "Generado con Encuesta Clara · Observatorio de Inteligencia Artificial — UCCuyo · "
        "observatorioia@uccuyo.edu.ar."
    )
    # Ajustar a ~200 palabras
    words = re.findall(r"\S+", text)
    if len(words) > 220:
        text = " ".join(words[:200]).rstrip(",.;:") + "."
    return text


# Compatibilidad: portada antigua (si algún script aún la llama)
def add_cover_band(
    doc: Document,
    title_main: str,
    subtitle: str,
    n_rows: int = 0,
    n_vars: int = 0,
    source_name: str = "",
    *,
    kicker: str = "INFORME INSTITUCIONAL",
) -> None:
    _ = (n_rows, n_vars, source_name)
    add_aesthetic_cover(doc, kicker=kicker, title=title_main, subtitle=subtitle)