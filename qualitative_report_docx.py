"""
Exportación del informe cualitativo integrado a Microsoft Word (.docx).
Formato académico formal: Times New Roman, interlineado 1,5, márgenes estándar.
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor


def _set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in ((1, 16), (2, 14), (3, 12)):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.name = "Times New Roman"
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(12 if level > 1 else 18)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def _add_runs_from_inline(paragraph: Any, text: str) -> None:
    """Interpreta **negrita**, *cursiva* y «comillas» en un párrafo."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|«[^»]+»|_[^_]+_|[^*_«»]+)",
        re.DOTALL,
    )
    for chunk in pattern.findall(text):
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("«") and chunk.endswith("»"):
            paragraph.add_run(chunk)
        elif chunk.startswith("_") and chunk.endswith("_"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        else:
            paragraph.add_run(chunk)


def _add_paragraph(doc: Document, text: str, *, style: str | None = None, quote: bool = False) -> None:
    clean = text.strip()
    if not clean:
        return
    if quote:
        p = doc.add_paragraph(style="Intense Quote")
        _add_runs_from_inline(p, clean.strip('"\''))
        return
    p = doc.add_paragraph(style=style)
    _add_runs_from_inline(p, clean)


def _parse_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def _is_table_separator(line: str) -> bool:
    s = line.strip().replace("|", "").replace(":", "").replace("-", "").replace(" ", "")
    return len(s) == 0 and "---" in line


def _add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs_from_inline(p, cell_text)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def markdown_to_docx_bytes(markdown: str, *, title: str = "Informe de análisis cualitativo") -> bytes:
    """Convierte Markdown del informe cualitativo a bytes .docx."""
    doc = Document()
    _set_document_style(doc)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "Times New Roman"
    doc.add_paragraph()

    lines = markdown.splitlines()
    i = 0
    bullet_buffer: list[str] = []
    table_buffer: list[str] = []

    def flush_bullets() -> None:
        nonlocal bullet_buffer
        for item in bullet_buffer:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_from_inline(p, item)
        bullet_buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [_parse_table_row(r) for r in table_buffer if not _is_table_separator(r)]
        _add_markdown_table(doc, rows)
        table_buffer = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            flush_bullets()
            flush_table()
            i += 1
            continue

        if line.strip() == "---":
            flush_bullets()
            flush_table()
            doc.add_paragraph()
            i += 1
            continue

        if line.lstrip().startswith("|"):
            flush_bullets()
            table_buffer.append(line)
            i += 1
            continue

        flush_table()

        if line.lstrip().startswith("# "):
            flush_bullets()
            doc.add_heading(line.lstrip()[2:].strip(), level=1)
            i += 1
            continue
        if line.lstrip().startswith("## "):
            flush_bullets()
            doc.add_heading(line.lstrip()[3:].strip(), level=2)
            i += 1
            continue
        if line.lstrip().startswith("### "):
            flush_bullets()
            doc.add_heading(line.lstrip()[4:].strip(), level=3)
            i += 1
            continue

        if line.lstrip().startswith("> "):
            flush_bullets()
            quote_text = line.lstrip()[2:].strip()
            _add_paragraph(doc, quote_text, quote=True)
            i += 1
            continue

        if line.lstrip().startswith(("- ", "* ")):
            bullet_buffer.append(line.lstrip()[2:].strip())
            i += 1
            continue

        if line.strip().startswith("_") and line.strip().endswith("_"):
            flush_bullets()
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("_"))
            run.italic = True
            run.font.size = Pt(11)
            i += 1
            continue

        flush_bullets()
        _add_paragraph(doc, line)
        i += 1

    flush_bullets()
    flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_qualitative_report_docx(
    markdown: str,
    *,
    question_label: str = "",
) -> bytes:
    """Genera el archivo Word del informe cualitativo integrado."""
    _ = question_label  # reservado para metadatos futuros en el documento
    return markdown_to_docx_bytes(markdown, title="Informe de análisis cualitativo")
