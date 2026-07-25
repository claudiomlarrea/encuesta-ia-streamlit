"""
Informe EJECUTIVO (Word) — narrativo.

Portada + índice clicable (con subpuntos de recomendaciones) + numeración de páginas.
Si se pasan cruces del Análisis automático, se agrega un apartado con lectura y tablas.
"""
from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

from executive_narrative import build_executive_sections
from report_common import (
    GREEN,
    MUTED,
    add_aesthetic_cover,
    add_heading,
    add_pandas_table,
    add_para,
    add_toc_hyperlink,
    narrative_for_crosstab,
    normalize_user_crosses,
    setup_margins,
)
from survey_intel import classify_columns


def _infer_audience(source_name: str, profiles) -> str:
    blob = " ".join([source_name or ""] + [p.name for p in profiles[:20]]).lower()
    if "docente" in blob:
        return "docentes"
    if "alumno" in blob or "estudiante" in blob or "unidad academica estás cursando" in blob:
        return "estudiantes"
    return "la comunidad universitaria"


def _rec_bookmark(sec_idx: int, rec_n: int) -> str:
    return f"bm_ex_{sec_idx}_r{rec_n}"


def build_executive_report_docx(
    df: pd.DataFrame,
    *,
    title: str = "",
    subtitle: str = "",
    cohort_label: str = "",
    source_name: str = "",
    user_crosses: list[dict[str, Any]] | None = None,
    auto_analysis: list[dict[str, Any]] | None = None,
) -> bytes:
    import report_common as rc

    rc._bookmark_seq = 0

    profiles = classify_columns(df)
    audience = _infer_audience(source_name, profiles)
    cover_title = title.strip() or (
        f"Diagnóstico institucional sobre el uso de Inteligencia Artificial en {audience} "
        "de la Universidad Católica de Cuyo"
    )
    cover_sub = subtitle.strip() or (
        "Documento de síntesis para presentación a la comunidad educativa de la UCCuyo."
    )

    if user_crosses is not None:
        crosses = list(user_crosses)
    elif auto_analysis is not None:
        crosses = normalize_user_crosses(auto_analysis)
    else:
        crosses = []

    sections = build_executive_sections(df, audience=audience)
    section_items = list(sections.items())

    # Subpuntos: recomendaciones numeradas del apartado 7 (si existen)
    rec_subpoints: dict[int, list[tuple[int, str]]] = {}
    for idx, (name, paras) in enumerate(section_items, start=1):
        subs: list[tuple[int, str]] = []
        for para in paras:
            m = re.match(r"^(\d+)\.\s+(.+)$", para.strip())
            if m and len(m.group(2)) > 12:
                subs.append((int(m.group(1)), para.strip()))
        if subs and ("recomend" in name.lower() or idx == 7):
            rec_subpoints[idx] = subs

    cross_sec_idx = len(section_items) + 1 if crosses else None

    doc = Document()
    add_aesthetic_cover(
        doc,
        kicker="INFORME EJECUTIVO INSTITUCIONAL",
        title=cover_title,
        subtitle=cover_sub,
    )
    setup_margins(doc, section=doc.sections[-1])

    add_heading(doc, "INFORME EJECUTIVO", 1, bookmark="bm_ex_titulo")
    add_para(
        doc,
        f"Uso de Inteligencia Artificial en {audience} de la Universidad Católica de Cuyo",
        size=13,
        bold=True,
        space_after=10,
    )
    if cohort_label:
        add_para(doc, f"Cohorte / instrumento: {cohort_label}.", size=10, color=MUTED)

    add_heading(doc, "ÍNDICE", 1, bookmark="bm_ex_indice")
    add_para(
        doc,
        "Hacé clic en cada ítem para ir al apartado correspondiente.",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=6,
    )
    for idx, (name, _paras) in enumerate(section_items, start=1):
        add_toc_hyperlink(doc, name, f"bm_ex_{idx}", bold=True, space_after=2)
        for rec_n, rec_text in rec_subpoints.get(idx, []):
            # acortar para el índice
            short = rec_text if len(rec_text) <= 110 else rec_text[:107].rstrip() + "…"
            add_toc_hyperlink(
                doc,
                short,
                _rec_bookmark(idx, rec_n),
                size=10,
                indent_cm=0.6,
                space_after=1,
            )

    if cross_sec_idx is not None:
        add_toc_hyperlink(
            doc,
            "Cruces seleccionados en el análisis automático",
            f"bm_ex_{cross_sec_idx}",
            bold=True,
            space_after=2,
        )
        for i, cr in enumerate(crosses, start=1):
            add_toc_hyperlink(
                doc,
                f"{i}. {cr['row_label']} × {cr['col_label']}",
                f"bm_ex_{cross_sec_idx}_{i}",
                size=10,
                indent_cm=0.6,
                space_after=1,
            )

    for idx, (name, paras) in enumerate(section_items, start=1):
        add_heading(doc, name, 1, bookmark=f"bm_ex_{idx}")
        for para in paras:
            m = re.match(r"^(\d+)\.\s+(.+)$", para.strip())
            if m and idx in rec_subpoints:
                rec_n = int(m.group(1))
                add_heading(doc, para.strip(), 2, bookmark=_rec_bookmark(idx, rec_n))
            elif m and para[1] == ".":
                add_para(doc, para, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
            else:
                add_para(doc, para)

    if cross_sec_idx is not None and crosses:
        add_heading(
            doc,
            "Cruces seleccionados en el análisis automático",
            1,
            bookmark=f"bm_ex_{cross_sec_idx}",
        )
        add_para(
            doc,
            "Este apartado incorpora los cruces configurados en el Análisis automático. "
            "Cada cruce se presenta con una lectura narrativa breve y la tabla de contingencia "
            "correspondiente, para que el documento ejecutivo refleje las mismas exploraciones "
            "bivariadas revisadas en la aplicación.",
        )
        for i, cr in enumerate(crosses, start=1):
            add_heading(
                doc,
                f"{i}. {cr['row_label']} × {cr['col_label']}",
                2,
                bookmark=f"bm_ex_{cross_sec_idx}_{i}",
            )
            add_para(doc, narrative_for_crosstab(cr))
            add_para(doc, "Tabla de contingencia", size=11, bold=True, color=GREEN, space_after=2)
            add_pandas_table(doc, cr["table"])
            if cr.get("conclusion"):
                add_para(doc, cr["conclusion"], size=10, color=MUTED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
